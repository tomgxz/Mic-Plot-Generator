from docx import Document
from docx.shared import Pt,Cm,Mm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.oxml.shared import OxmlElement

from msoffice2pdf import convert
import pythoncom,os

from .models import Mic,MicPos,Scene
from .utils import verifyShow, verifyAct

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_vertical_alignment(cell, align="center"): 
    try:   
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')  
        tcValign.set(qn('w:val'), align)  
        tcPr.append(tcValign)
        return True 
    except Exception as e:
        print(e)           
        return False


def getMicField(mic:Mic,sortby:int) -> str:
    return str(mic.packnumber if sortby > 0 else mic.mixchannel)

def mpHasActor(mp:MicPos) -> bool:
    return mp.actor not in [None,""]

def mpsHaveSameActor(x:MicPos,y:MicPos) -> bool:
    return (x.actor == y.actor)


def generateParamaters(showdict:dict,actdict:dict,sortby=0,onlyshowchanges=1):
    assert sortby in [0,1]
    assert onlyshowchanges in [0,1]

    """

    sortby=0 -> mic mixer channel
    sortby=1 -> mic pack number

    only show changes defines whether
    text is shown in every live cell,
    or only when an actor gains a new mic
    
    """

    sortbytext = "packnumber" if sortby > 0 else "mixchannel"

    # get mic count
    maxn = len(Mic.objects.filter(show=showdict["original"]))

    # make a blank row template based on the mic count + 1 for the first column
    blankrow = [None for _ in range(maxn+1)]
    starting = blankrow.copy()

    # create lists and dicts
    scenes = []
    micpos = []
    micmap = {}

    # get every mic sorted by packnumber or mixchannel (according to sortby)
    micssorted = Mic.objects.filter(show=showdict["original"]).order_by(sortbytext)

    """
    create a map of packnumber/mixchannel to column number
    so that you can query the packnumber/mixchannel in the
    micmap dict to get the column number
    """
    
    for mic in enumerate(micssorted):
        micfield = getMicField(mic[1],sortby)
        micmap[micfield] = mic[0]

    # for every mic, add all child micpos' in this act to micsorted
    for mic in micssorted:
        for mp in MicPos.objects.all().filter(mic=mic):
            if mp.mic.show == showdict["original"] and mp.scene.act == actdict["original"]:
                micpos.append(mp)

    # for every micpos
    for mp in micpos:
        micfield = getMicField(mp.mic,sortby)
        if micmap[micfield] <= len(starting):
            if starting[micmap[micfield]] is not None and starting[micmap[micfield]]["micpos"] is not None:
                if mp.scene.number < starting[micmap[micfield]]["micpos"].scene.number:
                    if len(mp.actor) > 0: starting[micmap[micfield]] = {"text":mp.actor,"speaking":mp.speaking,"micpos":mp}
            else: 
                if len(mp.actor) > 0: starting[micmap[micfield]] = {"text":mp.actor,"speaking":mp.speaking,"micpos":mp}

    # if every actor name is shown
    if onlyshowchanges == 0:

        for scene in enumerate(Scene.objects.filter(act=actdict["original"]).order_by("number")):
            scenes.append(blankrow.copy())
            scenes[scene[0]][0] = scene[1].number

            for mp in micpos:
                if mp.scene == scene[1] and mp.actor not in [None,""]:
                    micfield = getMicField(mp.mic,sortby)
                    scenes[scene[0]][micmap[micfield]+1] = {"text":mp.actor,"speaking":mp.speaking,"micpos":mp}

    # if only the mic changes are shown
    else:

        # set previous scene and scene count
        previousscene = None
        
        scenequery = Scene.objects.filter(act=actdict["original"]).order_by("number")
        scenecount = len(scenequery)

        # for every scene in the given act
        for scene in enumerate(scenequery):

            # create a new scene in the sceens dictionary
            scenes.append(blankrow.copy())

            # set the scene number in the first column
            scenes[scene[0]][0] = scene[1].number

            # for every micpos that matches the current scene and has an actor
            for mp in micpos:
                if mp.scene == scene[1] and mp.actor not in [None,""]:
                    # set the table data to be speaking, with no text
                    micfield = getMicField(mp.mic,sortby)
                    scenes[scene[0]][micmap[micfield]+1] = {"text":"","speaking":2,"micpos":mp}

            # if there is a previous scene
            if previousscene is not None:

                # for every mic sorted by the given filter
                for mic in micssorted:
                    """
                    CELL CONTAINS TEXT IF EITHER OF
                        this cell has an actor
                            and the previous cell has a different actor
                        this cell has no actor
                            and the previous cell has an actor
                            and the next mp with an actor has a different actor
                    """

                    # get the mic field
                    micfield = getMicField(mic,sortby)

                    # the cellContainsText var defines whether the table
                    # cell should show any actor text. It is set to true
                    # if the conditions defined above are met

                    # cellContent contains the content that the cell should
                    # have. It is currently set as an abstract dictionary
                    # that is defined later when cellContainsText is true

                    cellContainsText = False
                    cellContent = {"text":"","speaking":0,"micpos":None}

                    micposquery = MicPos.objects.filter(mic=mic,scene=scene[1])

                    # if the query has no matches, move on to next mic
                    if len(micposquery) == 0: continue
                    mp = micposquery[0]

                    """
                    this cell has an actor
                    and this is not the first scene
                    and the previous cell has a different actor
                    """
                    if mpHasActor(mp):
                        # get the previous micpos
                        previousmicposquery = MicPos.objects.filter(mic=mic,scene=previousscene)

                        # only continue if the query has any matches
                        if not len(previousmicposquery) == 0:
                            previousmicpos = previousmicposquery[0]

                            # only continue if the previous mp has an actor
                            if mpHasActor(previousmicpos):
                                
                                # only continue if the mps have different actors
                                if mpsHaveSameActor(mp,previousmicpos): 

                                    # if the mps have different actors, the current cell should contain text
                                    cellContainsText = True
                                    cellContent = {"text":mp.actor,"speaking":mp.speaking,"micpos":mp}

                    """
                    this cell has no actor
                    and the previous cell has an actor
                    and the next mp with an actor has a different actor
                    """
                    if not mpHasActor(mp):
                        # get the previous micpos
                        previousmicposquery = MicPos.objects.filter(mic=mic,scene=previousscene)

                        # only continue if the query has any matches
                        if not len(previousmicposquery) == 0:
                            previousmicpos = previousmicposquery[0]

                            # only continue if the previous mp has an actor
                            if mpHasActor(previousmicpos):

                                # get the next micpos with an actor

                                # set closest value to infinity as a starting point
                                closest = float("inf")

                                # for every MicPos in this mic
                                for next in MicPos.objects.filter(mic=mic):
                                    
                                    if (next.scene.number < closest and # the scene number is less than the previous closest value
                                        next.scene.number > scene[1].number and # the scene number is greater than the current scene
                                        mpHasActor(next)): # the micpos has an actor

                                        # set the closest value to this scene number
                                        closest = next.scene.number

                                # only continue if there is a closest value
                                if closest != float("inf"):

                                    nextscenequery = Scene.objects.filter(act=actdict["original"],number=closest)

                                    if len(nextscenequery) > 0:

                                        nextscene = nextscenequery[0]
                                        nextmicpos = MicPos.objects.get(mic=mic,scene=nextscene)

                                        # only continue if the next mic pos has an actor
                                        if mpHasActor(nextmicpos):
                                            
                                            # only continue if the mps have different actors
                                            if not mpsHaveSameActor(previousmicpos,nextmicpos):

                                                # if the two mps have different actors, the current cell should contain text
                                                cellContainsText = True
                                                cellContent = {"text":nextmicpos.actor,"speaking":0,"micpos":mp}

                    if cellContainsText:
                        scenes[scene[0]][micmap[micfield]+1] = cellContent


                    """

                    # get the next micpos that has actor text
                    closest = float("inf")
                    for next in MicPos.objects.filter(mic=mic):
                        if next.scene.number < closest and next.actor not in [None,""] and not(next.scene.number <= scene[1].number):
                            closest = next.scene.number

                    # if there is another MP
                    if closest != float("inf") and closest < scenecount:

                        # this block will get any matching MPs
                        for nextmp in MicPos.objects.filter(
                                mic=mic,
                                scene=Scene.objects.filter(act=actdict["original"],number=closest)[0]
                            ):

                            # if nextmp has an actor, and the actor is different to the previous scene's MP

                            oldmpactor = None
                            for oldmp in MicPos.objects.filter(mic=mic,scene=previousscene[1]):
                                oldmpactor = oldmp.actor

                            if nextmp.actor != oldmpactor and nextmp.actor not in [None,""]:
                                currentmp = MicPos.objects.filter(mic=mic,scene=scene[1])
                                if len(currentmp) > 0: 
                                    scenes[scene[0]][micmap[micfield]+1] = (f"{nextmp.actor}",mp.speaking,mp)
                                else: 
                                    scenes[scene[0]][micmap[micfield]+1] = f"{nextmp.actor}"

                            break
                            
                    """

            previousscene = scene

    records = [
        ["Char",*list(map(lambda x: str(x.packnumber if sortby > 0 else x.mixchannel),micssorted))],
        ["ACT 1",*list(map(lambda x: starting[micmap[str(x.packnumber if sortby > 0 else x.mixchannel)]]["text"] if starting[micmap[str(x.packnumber if sortby > 0 else x.mixchannel)]]["text"] is not None else "",micssorted))],
        *scenes,
    ]

    print(records)
 
    return records

def generateDocument(showdict):
    document = Document()

    style = document.styles["Normal"]
    font = style.font
    font.name = "Roboto"
    font.size = Pt(11)

    section = document.sections[0]

    # Set to A4 page size
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    # Set to landscape
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    
    # set footer
    footer = section.footer

    p = footer.paragraphs[0]
    p.text = f"MIC PLOT - {showdict['name']} - {showdict['date'].year}".upper()

    return document

def createMicPlotDocument(show_id,show_name):

    showdict = verifyShow(show_id,show_name,sortmicsby=0)
    document = generateDocument(showdict)

    # create two copies in the document, one that sorts by mixer channel (0)
    # and one that sorts by mic pack number (1)
    for sortby in [0,1]:

        # generate a copy for each act
        for actdict in showdict["acts"]:

            # get the table data for the act, based on the sortby value
            records = generateParamaters(showdict,actdict,sortby,onlyshowchanges=sortby)

            pagetitle = f"Mic plot for {actdict['name']} - SOUND DESK - "
            if sortby > 0: pagetitle = f"Mic plot for {actdict['name']} - "

            p = document.add_paragraph(pagetitle)

            # create table
            table = document.add_table(len(records),len(records[0]))
            table.style = 'Table Grid'
            table.autofit = True
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # create header row with thicker border
            for r in enumerate(records[0]):
                # create cell and set text
                c = table.cell(0,r[0])
                c.paragraphs[0].add_run(str(r[1])).bold=True

                # set font size of cell
                paragraph = c.paragraphs[0]
                runs = paragraph.runs
                for run in runs: run.font.size=Pt(7)

                borderformat = {"sz": 16, "val": "single", "color": "#000000"} 

                # set border to be thicker than others (16px)
                set_cell_border(
                    c,
                    top=borderformat,
                    bottom=borderformat,
                    start=borderformat,
                    end=borderformat,
                )

            # create rows for each scene (and starting mics)
            for row in enumerate(records):
                # ignore first row (generated in the header)
                if row[0] == 0: continue

                # iterate through cells
                for r in enumerate(row[1]):
                    c = table.cell(row[0],r[0])

                    # ensure the cell has content
                    if r[1] not in [None,""]:

                        """

                        # if it is a MicPos type (ie not a starting character cell)
                        if type(r[1]) == tuple:
                            # tuple syntax (actor:str,speaking:int,mp:MicPos=None)

                            # set cell content to actor
                            if sortby != 1 or True: 
                                c.text = str(r[1][0])

                            else:

                                closest = float("inf")
                                for next in MicPos.objects.filter(mic=r[1][2].mic):
                                    if next.scene.number <= r[1][2].scene.number: continue
                                    if next.scene.number < closest: closest = next.scene.number

                                try:
                                    next = MicPos.objects.filter(
                                        mic=r[1][2].mic,
                                        scene=Scene.objects.filter(act=actdict["original"],number=closest)[0]
                                    )[0]

                                    if next.actor != r[1][2].actor and next.actor not in [None,""]:
                                        c.text = f"TO {next.actor}"

                                except IndexError: # no micpos matching said query
                                    pass

                                except OverflowError: # closest still = infinity
                                    pass

                            # lookup next MicPos
                            # if actor is not this actor
                            # next cell needs to have that text

                            # set cell background based on speaking type
                            if r[1][1] > 0:
                                speakingxml = r'<w:shd {} w:fill="cccccc"/>'.format(nsdecls('w'))
                                if r[1][1] == 1: r'<w:shd {} w:fill="bcbcbc"/>'.format(nsdecls('w'))
                                c._tc.get_or_add_tcPr().append(parse_xml(speakingxml))

                        """

                        if type(r[1]) == dict:

                            c.text = str(r[1]["text"])

                            if r[1]["speaking"] > 0:
                                speakingxml = r'<w:shd {} w:fill="cccccc"/>'.format(nsdecls('w'))
                                if r[1]["speaking"] == 1: r'<w:shd {} w:fill="bcbcbc"/>'.format(nsdecls('w'))
                                c._tc.get_or_add_tcPr().append(parse_xml(speakingxml))

                        # if it is a string/int (ie a starting character cell)
                        else:
                            # set text to record content
                            if r[0] == 0: c.paragraphs[0].add_run(str(r[1])).bold=True
                            else: c.paragraphs[0].add_run(str(r[1]))

                        # set cell font size
                        paragraph = c.paragraphs[0]
                        runs = paragraph.runs
                        for run in runs: run.font.size=Pt(7)

                    borderformat = {"sz": 6, "val": "single", "color": "#000000"}

                    # set thin border (6px)
                    set_cell_border(
                        c,
                        top=borderformat,
                        bottom=borderformat,
                        start=borderformat,
                        end=borderformat,
                    )

            # set columns width to fit content & page
            for column in table.columns:
                for cell in column.cells:
                    cell._tc.tcPr.tcW.type = 'auto'
                    set_cell_vertical_alignment(cell)

            # set row height
            for row in table.rows: row.height = Cm(.65)

            # add page break if not last act
            if actdict != showdict["acts"][-1]: document.add_page_break()
        
        # add page break if not last sort function
        if sortby != 1: document.add_page_break()

    return document

def fileToPDF(path):
    output = convert(source=os.path.abspath("demo.docx"), output_dir=os.path.abspath("."), soft=0)
    if os.path.isfile("demo.pdf"): os.remove("demo.pdf")
    os.rename(output,"demo.pdf")

def saveAsDOCX(document):
    document.save("demo.docx")
    return "demo.docx"

pythoncom.CoInitialize()

